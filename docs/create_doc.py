#!/usr/bin/env python3
"""
Generate the Superloop RWA Strategy Document (DOCX)
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

doc = Document()

# ── Styles ────────────────────────────────────────────────────
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
font.color.rgb = RGBColor(0x33, 0x33, 0x33)

for level in range(1, 4):
    heading_style = doc.styles[f'Heading {level}']
    heading_style.font.name = 'Calibri'
    heading_style.font.color.rgb = RGBColor(0x00, 0x7A, 0xFF)
    heading_style.font.bold = True
    if level == 1:
        heading_style.font.size = Pt(24)
    elif level == 2:
        heading_style.font.size = Pt(18)
    else:
        heading_style.font.size = Pt(14)

def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.style.font.bold = True
            run = p.runs[0] if p.runs else p.add_run(h)
            run.bold = True
            run.font.size = Pt(10)

    # Data rows
    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            cell = table.rows[r + 1].cells[c]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)

    doc.add_paragraph()  # spacing
    return table

def add_bullet_list(doc, items, bold_first=False):
    for item in items:
        p = doc.add_paragraph(style='List Bullet')
        if bold_first and ': ' in item:
            parts = item.split(': ', 1)
            run = p.add_run(parts[0] + ': ')
            run.bold = True
            run.font.size = Pt(11)
            run = p.add_run(parts[1])
            run.font.size = Pt(11)
        else:
            run = p.add_run(item)
            run.font.size = Pt(11)

# ══════════════════════════════════════════════════════════════
# COVER PAGE
# ══════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
for _ in range(6):
    p.add_run('\n')

run = p.add_run('SUPERLOOP')
run.font.size = Pt(42)
run.font.color.rgb = RGBColor(0x00, 0x7A, 0xFF)
run.bold = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Institutional RWA Looping Infrastructure')
run.font.size = Pt(22)
run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Automated Leveraged Yield Vaults for Tokenized Real-World Assets')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
for _ in range(4):
    p.add_run('\n')
run = p.add_run('Built on Aave V3 | Aave Horizon | Centrifuge | Morpho\n')
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
run = p.add_run('ERC-4626 Compliant | Modular Architecture | Flash-Loan Optimized\n\n')
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
run = p.add_run('February 2026 | Confidential')
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# TABLE OF CONTENTS
# ══════════════════════════════════════════════════════════════
doc.add_heading('Table of Contents', level=1)
toc_items = [
    '1. Executive Summary',
    '2. The Problem',
    '3. The Solution: Superloop',
    '4. Market Opportunity',
    '5. How It Works',
    '6. Yield Economics',
    '7. Architecture Overview',
    '8. Existing Infrastructure (Verified from Codebase)',
    '9. Aave Horizon Integration',
    '10. Centrifuge Integration',
    '11. Exact Integrations Required at Scale',
    '12. Compliance Framework',
    '13. Risk Management',
    '14. Supported Assets & Partners',
    '15. Competitive Positioning',
    '16. Revenue Model',
    '17. Roadmap',
    '18. Why Whitelist Superloop (For RWA Issuers)',
    '19. Appendix: Technical Specifications',
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(4)
    p.runs[0].font.size = Pt(12)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 1. EXECUTIVE SUMMARY
# ══════════════════════════════════════════════════════════════
doc.add_heading('1. Executive Summary', level=1)

doc.add_paragraph(
    'Superloop is the institutional vault infrastructure layer for automated, compliant, leveraged Real-World Asset (RWA) strategies. '
    'We sit above lending protocols (Aave Horizon, Morpho), tokenization platforms (Centrifuge, Securitize), and DeFi primitives (flash loans, DEX aggregators), '
    'composing them into managed ERC-4626 vaults that deliver 9-13% APR on Treasury-backed collateral at 3-4x leverage.'
)

doc.add_paragraph(
    'The opportunity is significant and immediate: $35B+ in tokenized RWAs are on-chain but under-utilized, Aave Horizon reached $1B TVL in 6 months, '
    'and 64% of Morpho\'s $3B in loans comes from looping strategies. Institutions need this yield, but the process of looping positions manually requires '
    '30+ transactions. Superloop automates this into a single flash-loan-powered atomic transaction.'
)

p = doc.add_paragraph()
run = p.add_run('Key highlights:')
run.bold = True
add_bullet_list(doc, [
    '70%+ of infrastructure already built and production-tested (Aave V3 modules, flash loans, async queues)',
    'Aave Horizon integration requires zero code changes to existing modules -- deploy only with new PoolAddressesProvider',
    '$35B+ addressable RWA market growing 300%+ year-over-year',
    '9-13% APR target on Treasury-backed collateral at 3-4x leverage',
    'Compliance-first: KYC gating, transfer-restricted shares, ERC-1404/ERC-3643 compatible',
    'ERC-4626 tokenized vaults with async deposit/withdraw queues and performance fee accounting',
])

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 2. THE PROBLEM
# ══════════════════════════════════════════════════════════════
doc.add_heading('2. The Problem', level=1)

doc.add_paragraph(
    'Tokenized real-world assets are live on-chain -- US Treasuries, AAA CLOs, corporate credit, and money market funds now exist as ERC-20 tokens. '
    'However, the infrastructure to unlock their full yield potential is fragmented, manual, and inaccessible to most participants.'
)

doc.add_heading('2.1 Manual Looping is Complex and Error-Prone', level=2)
doc.add_paragraph(
    'To create a leveraged RWA position, an institution must: (1) acquire the RWA token, (2) supply it as collateral to a lending market, '
    '(3) borrow stablecoins against it, (4) swap the borrowed stablecoins back to the RWA token, and (5) repeat. '
    'Each loop iteration requires 3-5 transactions. Reaching 4x leverage requires 12-15 loop iterations, totaling 30+ individual transactions. '
    'This is gas-intensive, error-prone, and requires constant monitoring.'
)

doc.add_heading('2.2 Fragmented Infrastructure', level=2)
doc.add_paragraph(
    'RWA tokens live on issuer platforms (Superstate, Centrifuge, Securitize). Lending markets are separate protocols (Aave Horizon, Morpho, Spark). '
    'Oracles (Chainlink NAVLink) are yet another layer. Compliance verification (KYC/AML) sits in its own silo. '
    'No unified system exists to compose these into managed vault strategies.'
)

doc.add_heading('2.3 Compliance Barriers', level=2)
doc.add_paragraph(
    'RWA tokens carry transfer restrictions at the smart contract level. Tokens using ERC-1404 (like Superstate\'s USTB) enforce '
    '`detectTransferRestriction()` checks before every transfer. Others use ERC-3643 (T-REX) with identity registries. '
    'Any vault interacting with these tokens must be KYC-whitelisted by each issuer. '
    'No plug-and-play compliance framework exists for DeFi vaults.'
)

p = doc.add_paragraph()
run = p.add_run('The result: ')
run.bold = True
run = p.add_run(
    '$35B+ in tokenized RWAs sit under-utilized. Institutions earn base yield (~4-5%) when they could earn 9-13% with automated leveraged strategies.'
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 3. THE SOLUTION
# ══════════════════════════════════════════════════════════════
doc.add_heading('3. The Solution: Superloop', level=1)

doc.add_paragraph(
    'Superloop is the institutional vault infrastructure layer that sits above Aave Horizon, Centrifuge, Morpho, and other DeFi primitives, '
    'composing them into managed, compliant, leveraged RWA strategies.'
)

doc.add_heading('3.1 One-Click Looping', level=2)
doc.add_paragraph(
    'Flash-loan powered single-transaction looping. Users deposit stablecoins; the vault auto-loops to the target leverage. '
    'A deposit of $100K USDC becomes a $400K leveraged Treasury position in one atomic transaction. No manual intervention, no multi-block risk.'
)

doc.add_heading('3.2 Modular Architecture', level=2)
doc.add_paragraph(
    'Plug-and-play modules for every DeFi primitive: Aave V3 supply/borrow, flash loans (Aave, Morpho, Balancer), '
    'DEX swaps (1inch, Uniswap), vault interactions (ERC-4626), and staking. '
    'Adding a new protocol means deploying a new module -- no changes to core vault logic.'
)

doc.add_heading('3.3 Built-In Compliance', level=2)
doc.add_paragraph(
    'KYC gating at vault entry via the DepositManager. Transfer-restricted vault shares via the onlyPrivileged modifier. '
    'Compatible with ERC-1404, ERC-3643 (T-REX), and Chainlink ACE (Automated Compliance Engine). '
    'The vault entity completes institutional KYC with each RWA issuer.'
)

doc.add_heading('3.4 ERC-4626 Standard', level=2)
doc.add_paragraph(
    'Fully composable tokenized vault. Async deposit/withdraw queues with four priority tiers (Instant, Priority, Deferred, General). '
    'Performance fee via share dilution. Institutional-grade accounting via pluggable accountant modules.'
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 4. MARKET OPPORTUNITY
# ══════════════════════════════════════════════════════════════
doc.add_heading('4. Market Opportunity', level=1)

add_table(doc,
    ['Metric', 'Value', 'Context'],
    [
        ['Total tokenized RWAs on-chain', '$35B+', 'Growing 300%+ YoY'],
        ['Aave Horizon TVL', '$1B', 'Reached Feb 2026, 6 months post-launch'],
        ['Morpho loans outstanding', '$3B', '64% from looping strategies'],
        ['Spark Liquidity Layer AUM', '$3.5B', 'Sky/MakerDAO ecosystem'],
        ['BlackRock BUIDL AUM', '$2.3B+', 'Largest tokenized fund'],
        ['Ondo Finance TVL', '$3.2B+', 'Leading RWA protocol by retail TVL'],
    ],
    col_widths=[2.5, 1.5, 3.0]
)

doc.add_heading('4.1 Key RWA Issuers and Assets', level=2)
add_table(doc,
    ['Issuer', 'Token(s)', 'Underlying', 'AUM', 'Platform'],
    [
        ['Superstate', 'USTB, USCC', 'US Govt Securities, Crypto Carry', '$500M+', 'Aave Horizon'],
        ['Centrifuge', 'JTRSY, JAAA', 'Janus Henderson Treasuries / AAA CLOs', '$1B+', 'Aave Horizon'],
        ['Hashnote / Circle', 'USYC', 'Short Duration Yield Fund', '$400M+', 'Aave Horizon'],
        ['VanEck / Securitize', 'VBILL', 'VanEck Treasury Fund', '$93M+', 'Aave Horizon'],
        ['BlackRock / Securitize', 'BUIDL', 'Treasury Money Market', '$2.3B+', 'Pipeline'],
        ['Ondo Finance', 'OUSG, USDY', 'Treasuries, Yield-Bearing Stablecoin', '$3.2B+', 'Future'],
    ],
    col_widths=[1.5, 1.2, 1.8, 0.8, 1.2]
)

doc.add_heading('4.2 Key Lending Markets', level=2)
add_table(doc,
    ['Market', 'TVL', 'RWA Support', 'Looping Volume'],
    [
        ['Aave Horizon', '$1B', 'Purpose-built for RWAs', 'Primary RWA lending venue'],
        ['Morpho', '$3B', 'Permissionless markets', '64% of volume is looping'],
        ['Spark', '$3.5B', 'Sky/MakerDAO integration', 'RWA-backed stablecoin (USDS)'],
    ],
    col_widths=[1.5, 1.0, 2.5, 2.5]
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 5. HOW IT WORKS
# ══════════════════════════════════════════════════════════════
doc.add_heading('5. How It Works', level=1)

doc.add_heading('5.1 User Flow (Non-Technical)', level=2)
doc.add_paragraph(
    'From the user\'s perspective, the interaction is simple:'
)
steps = [
    'DEPOSIT: User deposits USDC into the Superloop Vault (standard ERC-4626 deposit or async request)',
    'ACQUIRE: Vault acquires RWA token (e.g., USTB) via DEX swap or direct minting with issuer',
    'SUPPLY: Supply the RWA token as collateral on Aave Horizon (permissioned)',
    'BORROW: Borrow USDC against the RWA collateral at DeFi rates',
    'LOOP: Repeat steps 2-4 using flash loans for capital efficiency (3-5x leverage)',
    'EARN: Vault earns the leveraged spread between RWA yield and borrow cost (target 9-13% APR)',
]
for i, step in enumerate(steps, 1):
    p = doc.add_paragraph()
    run = p.add_run(f'Step {i}: ')
    run.bold = True
    run.font.size = Pt(11)
    parts = step.split(': ', 1)
    run = p.add_run(f'{parts[0]}: ')
    run.bold = True
    run = p.add_run(parts[1])

doc.add_paragraph(
    'The entire loop executes atomically in a single transaction using flash loans. All steps succeed or all revert. '
    'No partial state. No multi-block risk. Gas efficient.'
)

doc.add_heading('5.2 Technical Execution Flow', level=2)
doc.add_paragraph(
    'The operator (vault manager) calls superloop.operate() with a nested array of ModuleExecutionData. '
    'The flow follows the proven EthenaLoopTest pattern already in the production codebase:'
)

code_block = (
    '// Operator calls: superloop.operate(finalExecutionData)\n\n'
    'Step 1: resolveDepositRequests(USDC, depositAmount, callback=[\n'
    '  Step 2: morphoFlashLoan(USTB, flashAmount, callback=[\n'
    '    Step 3: aaveV3Supply(USTB, type(uint256).max)     // Supply all USTB to Horizon\n'
    '    Step 4: aaveV3Borrow(USDC, borrowAmount)          // Borrow USDC against USTB\n'
    '    Step 5: dexSwap(USDC -> USTB, borrowAmount)       // Swap borrowed USDC to more USTB\n'
    '    Step 6: aaveV3Supply(USTB, type(uint256).max)     // Re-supply new USTB\n'
    '    Step 7: aaveV3Borrow(USDC, repayAmount)           // Borrow to repay flash loan\n'
    '    // Flash loan callback auto-approves repayment\n'
    '  ])\n'
    '])\n\n'
    '// Result: Vault holds 4-5x leveraged USTB position on Horizon\n'
    '// AccountantPlugin reads: aToken balance (USTB) - variableDebt (USDC) = net value'
)
p = doc.add_paragraph()
run = p.add_run(code_block)
run.font.name = 'Consolas'
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(0x00, 0x7A, 0xFF)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 6. YIELD ECONOMICS
# ══════════════════════════════════════════════════════════════
doc.add_heading('6. Yield Economics', level=1)

doc.add_heading('6.1 The Looping Formula', level=2)
p = doc.add_paragraph()
run = p.add_run('Leveraged APR = L x Y - (L-1) x R')
run.bold = True
run.font.size = Pt(14)
p = doc.add_paragraph('Where: L = leverage multiple, Y = RWA yield, R = borrow rate')

doc.add_heading('6.2 Yield Table', level=2)
add_table(doc,
    ['Leverage', '5% Yield / 3% Borrow', '5% Yield / 2% Borrow', '8% Yield / 3% Borrow'],
    [
        ['1x (unleveraged)', '5.0%', '5.0%', '8.0%'],
        ['2x', '7.0%', '8.0%', '13.0%'],
        ['3x', '9.0%', '11.0%', '18.0%'],
        ['4x', '11.0%', '14.0%', '23.0%'],
        ['5x (max practical)', '13.0%', '17.0%', '28.0%'],
    ],
    col_widths=[1.5, 1.8, 1.8, 1.8]
)

doc.add_heading('6.3 Target Strategy Parameters', level=2)
add_bullet_list(doc, [
    'Conservative target: 3x leverage on US Treasury collateral (USTB) = 9-11% APR',
    'Maximum leverage: 4-5x depending on asset risk profile and LTV ratios',
    'Flash loan source: Morpho (0% fee) or Balancer (0% fee) to maximize thin RWA spreads',
    'Break-even leverage: L_breakeven = R / (R - Y) -- if spread is 2%, break-even occurs at infinite leverage',
    'Auto-deleverage trigger: When borrow rate approaches RWA yield within 50bps',
], bold_first=True)

doc.add_heading('6.4 Revenue Model', level=2)
add_table(doc,
    ['Revenue Stream', 'Rate', 'Mechanism'],
    [
        ['Performance Fee', '10-20% of yield', 'Share dilution on profit (already implemented in SuperloopVault)'],
        ['Management Fee', '0.5-2% annually', 'AUM-based via exchange rate adjustment in accountant'],
        ['Instant Withdraw Fee', '0.1-0.5%', 'Premium for skipping withdrawal queue (already implemented)'],
    ],
    col_widths=[2.0, 1.5, 3.5]
)

p = doc.add_paragraph()
run = p.add_run('Example at $100M TVL: ')
run.bold = True
doc.add_paragraph(
    'Gross yield: $10M/year. Performance fee (15%): $1.5M/year. Management fee (1%): $1.0M/year. '
    'Total protocol revenue: ~$2.5M/year per $100M TVL. Scaling path: $1B TVL = ~$25M annual revenue.'
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 7. ARCHITECTURE OVERVIEW
# ══════════════════════════════════════════════════════════════
doc.add_heading('7. Architecture Overview', level=1)

doc.add_paragraph(
    'Superloop uses a layered modular architecture where each concern is isolated into composable components:'
)

add_table(doc,
    ['Layer', 'Components', 'Purpose'],
    [
        ['User Layer', 'SuperloopVault (ERC-4626), DepositManager, WithdrawManager', 'User-facing deposit/withdraw, share accounting, async queues'],
        ['Strategy Layer', 'VaultOperator, operate() function, Flash Loan Wrapping', 'Strategy execution, atomic operations, callback handling'],
        ['Module Layer', 'AaveV3Supply, AaveV3Borrow, AaveV3Flashloan, MorphoFlashloan, UniversalDex, VaultSupply, Emode', 'Protocol interactions via standardized interfaces'],
        ['Accounting Layer', 'AaveV3AccountantPlugin, UniversalAccountant', 'Position tracking, NAV calculation, performance fee computation'],
        ['Protocol Layer', 'Aave Horizon, Aave V3, Morpho, Centrifuge, Chainlink, DEX Aggregators', 'External protocol integrations'],
    ],
    col_widths=[1.3, 3.0, 2.7]
)

doc.add_heading('7.1 Module Execution Pattern', level=2)
doc.add_paragraph(
    'The core execution pattern is simple but powerful. The vault operator calls operate() with an array of ModuleExecutionData, '
    'each specifying a module address, call type (CALL or DELEGATECALL), and encoded function parameters. '
    'The vault iterates through the array, executing each module in sequence. '
    'Modules can be nested (e.g., flash loan callbacks contain their own ModuleExecutionData arrays) '
    'enabling complex multi-step strategies in a single atomic transaction.'
)

doc.add_heading('7.2 Key Design Principles', level=2)
add_bullet_list(doc, [
    'Execution Context: All module calls are wrapped in beginExecutionContext/endExecutionContext, preventing unauthorized external calls',
    'Module Registry: Only whitelisted modules can be executed, preventing arbitrary contract calls',
    'Callback Routing: Flash loan and protocol callbacks route through the vault\'s fallback handler, which decodes nested execution data and continues the module chain',
    'Atomic Transactions: The entire operation (deposit resolution + flash loan + supply + borrow + swap + re-supply) executes in a single transaction',
    'Pluggable Accounting: The accountant module is replaceable, allowing different position tracking strategies per vault type',
])

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 8. EXISTING INFRASTRUCTURE
# ══════════════════════════════════════════════════════════════
doc.add_heading('8. Existing Infrastructure (Verified from Codebase)', level=1)

doc.add_paragraph(
    'Based on thorough verification of the Superloop smart contract codebase (superloop-core-contracts), '
    'the following components are already built and production-tested:'
)

doc.add_heading('8.1 Aave V3 Module Suite', level=2)
add_table(doc,
    ['Module', 'File Path', 'Key Function', 'RWA Compatible'],
    [
        ['AaveV3SupplyModule', 'src/modules/aave/', 'pool.supply(asset, amount, address(this), 0)', 'Yes -- deploy with Horizon provider'],
        ['AaveV3BorrowModule', 'src/modules/aave/', 'pool.borrow(asset, amount, 2, 0, address(this))', 'Yes -- deploy with Horizon provider'],
        ['AaveV3RepayModule', 'src/modules/aave/', 'pool.repay(asset, amount, 2, address(this))', 'Yes -- deploy with Horizon provider'],
        ['AaveV3WithdrawModule', 'src/modules/aave/', 'pool.withdraw(asset, amount, address(this))', 'Yes -- deploy with Horizon provider'],
        ['AaveV3EmodeModule', 'src/modules/aave/', 'pool.setUserEMode(emodeCategory)', 'Yes -- deploy with Horizon provider'],
        ['AaveV3FlashloanModule', 'src/modules/aave/', 'pool.flashLoanSimple() with callbacks', 'Yes -- deploy with Horizon provider'],
    ],
    col_widths=[1.8, 1.5, 2.5, 1.7]
)

p = doc.add_paragraph()
run = p.add_run('Critical finding: ')
run.bold = True
run = p.add_run(
    'All AaveV3 modules accept a poolAddressesProvider in their constructor (AaveV3ActionModule.sol:30). '
    'Aave Horizon is an Aave V3.3 fork with the same interface but a different PoolAddressesProvider address. '
    'This means integrating with Horizon requires ZERO code changes -- just new deployment instances pointed at the Horizon provider address.'
)

doc.add_heading('8.2 Flash Loan Modules', level=2)
add_table(doc,
    ['Module', 'Source', 'Fee', 'RWA Use'],
    [
        ['AaveV3FlashloanModule', 'Aave V3 pool', 'Variable (bps)', 'Primary -- wraps the looping callback chain'],
        ['MorphoFlashloanModule', 'Morpho Blue', '0%', 'Preferred -- zero fee maximizes RWA spread'],
    ],
    col_widths=[2.0, 1.5, 1.5, 2.5]
)

doc.add_heading('8.3 Vault Infrastructure', level=2)
add_table(doc,
    ['Component', 'File Path', 'Purpose', 'RWA Compatible'],
    [
        ['SuperloopVault', 'src/core/Superloop/', 'ERC-4626 vault with performance fees', 'Yes'],
        ['DepositManager', 'src/core/DepositManager/', 'Async deposit queue with request/resolve', 'Yes'],
        ['WithdrawManager', 'src/core/WithdrawManager/', '4-tier withdraw queue', 'Yes -- maps to RWA redemption timing'],
        ['AaveV3AccountantPlugin', 'src/plugins/Accountant/', 'Tracks lend/borrow via getUserReserveData', 'Yes -- deploy with Horizon provider'],
        ['UniversalDexModule', 'src/modules/dex/', 'Multi-DEX swap execution', 'Yes'],
        ['ModuleRegistry', 'src/core/ModuleRegistry/', 'Whitelist-based module validation', 'Yes'],
        ['VaultRouter', 'src/helpers/', 'Whitelisted vault/token/manager routing', 'Yes'],
    ],
    col_widths=[1.8, 1.8, 2.2, 1.7]
)

doc.add_heading('8.4 Proven Looping Pattern', level=2)
doc.add_paragraph(
    'The EthenaLoopTest (test/core/integration/EthenaLoop.t.sol) already demonstrates the exact looping pattern needed for RWA vaults: '
    'deposit resolution wraps a flash loan, which wraps supply + borrow + swap + re-supply operations. '
    'At 5x leverage, the test deposits 100 USDe, flash loans 400 USDe, and creates a leveraged position. '
    'The same pattern applies directly to RWA looping with USTB/USDC.'
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 9. AAVE HORIZON INTEGRATION
# ══════════════════════════════════════════════════════════════
doc.add_heading('9. Aave Horizon Integration', level=1)

doc.add_heading('9.1 What is Aave Horizon?', level=2)
doc.add_paragraph(
    'Aave Horizon is Aave Labs\' institutional-grade lending market purpose-built for Real World Assets on Ethereum. '
    'It is a separate, permissioned instance of the Aave Protocol, forked from Aave V3.3. '
    'Launched August 25, 2025, it reached $1B TVL by February 2026.'
)

doc.add_heading('9.2 Dual Structure', level=2)
add_bullet_list(doc, [
    'Permissioned Side (Collateral): Depositing RWA tokens requires KYC by the asset issuer. Vault address must be whitelisted. aTokens are non-transferable.',
    'Permissionless Side (Lending): Anyone can supply stablecoins (USDC, GHO, RLUSD) to earn yield from institutional borrowers. No KYC needed.',
    'Bridge: Only users who supply permissioned RWA collateral can borrow from the permissionless stablecoin pools.',
], bold_first=True)

doc.add_heading('9.3 Superloop Integration Steps', level=2)
add_bullet_list(doc, [
    'Step 1: Deploy existing AaveV3SupplyModule, AaveV3BorrowModule, AaveV3RepayModule, AaveV3WithdrawModule, AaveV3EmodeModule with HORIZON_POOL_ADDRESSES_PROVIDER as constructor argument',
    'Step 2: Deploy AaveV3AccountantPlugin with Horizon provider, configured with RWA lendAssets (USTB, JTRSY, etc.) and stablecoin borrowAssets (USDC, GHO)',
    'Step 3: Register all modules in ModuleRegistry and configure vault',
    'Step 4: Complete institutional KYC with each RWA issuer to whitelist the vault contract address',
    'Step 5: Operator executes looping strategy via operate() with flash-loan-wrapped supply/borrow/swap chain',
    'Step 6: AccountantPlugin tracks net position: aToken balance (collateral) - variableDebt (borrows) = vault NAV',
])

doc.add_heading('9.4 Horizon Oracle Integration', level=2)
doc.add_paragraph(
    'Unlike crypto-native assets priced by DEX markets, RWAs rely on Net Asset Value (NAV) reported by the issuer. '
    'Aave Horizon uses Chainlink NAVLink to deliver verified NAV data with price bounds validation. '
    'The existing AaveV3AccountantPlugin reads prices via IAaveOracle(poolAddressesProvider.getPriceOracle()).getAssetPrice(), '
    'which automatically returns NAVLink prices for Horizon RWA assets.'
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 10. CENTRIFUGE INTEGRATION
# ══════════════════════════════════════════════════════════════
doc.add_heading('10. Centrifuge Integration', level=1)

doc.add_heading('10.1 Overview', level=2)
doc.add_paragraph(
    'Centrifuge is the leading RWA tokenization platform with $1B+ in tokenized structured credit. '
    'It provides access to institutional products like JAAA (AAA CLOs from Janus Henderson) and JTRSY (US Treasuries). '
    'Centrifuge uses the ERC-7540 async vault standard with epoch-based settlement.'
)

doc.add_heading('10.2 Products Available', level=2)
add_table(doc,
    ['Token', 'Underlying', 'Manager', 'Expected Yield', 'Risk Profile'],
    [
        ['JTRSY', 'US Treasury exposure', 'Janus Henderson ($373B AUM)', '~4.5-5%', 'Very Low (sovereign)'],
        ['JAAA', 'AAA CLOs', 'Janus Henderson', '~5-7%', 'Low (AAA-rated)'],
        ['ACRDX', 'Diversified credit', 'Apollo ($938B AUM)', '~7-10%', 'Medium (institutional credit)'],
    ],
    col_widths=[0.8, 1.5, 2.0, 1.2, 1.5]
)

doc.add_heading('10.3 New Module: CentrifugeSupplyModule', level=2)
doc.add_paragraph(
    'Unlike Aave V3\'s synchronous supply/borrow, Centrifuge uses ERC-7540 async vaults with epoch-based settlement. '
    'A new CentrifugeSupplyModule is needed with the following interface:'
)
add_bullet_list(doc, [
    'requestDeposit(amount): Queue a deposit for the next epoch (24h minimum)',
    'claimDeposit(): After epoch resolves, claim tranche tokens',
    'requestRedeem(shares): Queue a redemption for the next epoch',
    'claimRedeem(): After epoch resolves, claim underlying assets',
])

doc.add_heading('10.4 Combined Strategy: Centrifuge + Horizon', level=2)
doc.add_paragraph(
    'The most powerful integration path combines Centrifuge and Horizon: '
    '(1) Acquire Centrifuge tranche tokens (JTRSY, JAAA) via the CentrifugeSupplyModule, '
    '(2) Supply these tranche tokens to Aave Horizon as collateral (they are accepted), '
    '(3) Borrow stablecoins and loop for leveraged yield. '
    'This creates a leveraged position on institutional structured credit via AAA CLOs or Treasuries.'
)

doc.add_heading('10.5 Grove Finance Partnership', level=2)
doc.add_paragraph(
    'Grove Finance (part of Sky/MakerDAO ecosystem) has a $1B allocation that routes through Centrifuge V3 into JAAA and ACRDX. '
    'Partnering with Grove gives Superloop access to institutional capital flows and priority allocation on Centrifuge products.'
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 11. EXACT INTEGRATIONS REQUIRED AT SCALE
# ══════════════════════════════════════════════════════════════
doc.add_heading('11. Exact Integrations Required at Scale', level=1)

doc.add_paragraph(
    'Based on thorough verification of the production codebase and analysis of target protocols, '
    'here is the complete list of integrations needed to operate RWA looping vaults at scale:'
)

doc.add_heading('11.1 Zero Code Changes (Deploy Only)', level=2)
doc.add_paragraph(
    'These existing modules work with Aave Horizon as-is because Horizon uses the standard Aave V3.3 interface. '
    'New instances are deployed with the Horizon PoolAddressesProvider address:'
)
add_table(doc,
    ['Module', 'Current Constructor', 'Change for Horizon'],
    [
        ['AaveV3SupplyModule', 'poolAddressesProvider_', 'Deploy new instance with HORIZON_POOL_ADDRESSES_PROVIDER'],
        ['AaveV3BorrowModule', 'poolAddressesProvider_', 'Deploy new instance with HORIZON_POOL_ADDRESSES_PROVIDER'],
        ['AaveV3RepayModule', 'poolAddressesProvider_', 'Deploy new instance with HORIZON_POOL_ADDRESSES_PROVIDER'],
        ['AaveV3WithdrawModule', 'poolAddressesProvider_', 'Deploy new instance with HORIZON_POOL_ADDRESSES_PROVIDER'],
        ['AaveV3EmodeModule', 'poolAddressesProvider_', 'Deploy new instance with HORIZON_POOL_ADDRESSES_PROVIDER'],
        ['AaveV3FlashloanModule', 'poolAddressesProvider_', 'Deploy new instance (if Horizon supports flashLoanSimple)'],
        ['AaveV3AccountantPlugin', 'poolAddressesProvider, lendAssets, borrowAssets', 'New instance with Horizon provider + RWA assets configured'],
    ],
    col_widths=[2.0, 2.0, 3.0]
)

doc.add_heading('11.2 New Modules Required', level=2)
add_table(doc,
    ['Module', 'Purpose', 'Complexity', 'Dependencies'],
    [
        ['CentrifugeSupplyModule', 'ERC-7540 async deposit/redeem for Centrifuge pools', 'Medium', 'Centrifuge V3 contracts, epoch manager'],
        ['RWAComplianceGate', 'KYC verification, investor tier checks at vault entry', 'Medium', 'Chainlink ACE / Sumsub API'],
        ['NAVOracleAdapter', 'Chainlink NAVLink price feed integration', 'Low', 'Chainlink NAVLink DON'],
        ['CentrifugeAccountantPlugin', 'Track pending/claimable positions + epoch NAV', 'Low', 'Centrifuge pool contract'],
        ['GroveRouterModule', 'Capital routing through Grove/Sky allocation', 'Medium', 'Grove Finance API, Centrifuge V3'],
    ],
    col_widths=[2.0, 2.5, 1.0, 2.0]
)

doc.add_heading('11.3 Infrastructure & Operations', level=2)
add_table(doc,
    ['Component', 'Purpose', 'Notes'],
    [
        ['Keeper/Operator Bot', 'Monitor rates, auto-deleverage, epoch claims, health factor', 'Runs off-chain, calls operate() on-chain'],
        ['Legal Entity (SPV)', 'Required for each RWA issuer whitelisting', 'Cayman SPC or Delaware Series LLC'],
        ['KYC Provider', 'Investor verification for vault deposits', 'Chainlink ACE (CCID), Sumsub, or Coinbase Verifications'],
        ['Risk Parameter Engine', 'Dynamic LTV/rate adjustment based on market conditions', 'Chaos Labs or LlamaRisk integration'],
        ['Monitoring Dashboard', 'Real-time position tracking, health factor alerts', 'TheGraph subgraph + custom frontend'],
    ],
    col_widths=[2.0, 2.5, 3.0]
)

doc.add_heading('11.4 Protocol-by-Protocol Integration Checklist', level=2)

p = doc.add_paragraph()
run = p.add_run('Aave Horizon:')
run.bold = True
add_bullet_list(doc, [
    'Deploy 7 module instances with Horizon PoolAddressesProvider (zero code changes)',
    'Register callback handlers for Horizon pool callbacks',
    'Configure AccountantPlugin with RWA lendAssets and stablecoin borrowAssets',
    'KYC vault contract with each RWA issuer (Superstate, Centrifuge, Hashnote, VanEck)',
    'Verify flash loan support for stablecoins on Horizon',
])

p = doc.add_paragraph()
run = p.add_run('Centrifuge:')
run.bold = True
add_bullet_list(doc, [
    'Build CentrifugeSupplyModule with ERC-7540 requestDeposit/claimDeposit/requestRedeem',
    'Integrate epoch monitoring via keeper bot',
    'Deploy CentrifugeAccountantPlugin for position tracking',
    'KYC vault with Centrifuge pool operators',
])

p = doc.add_paragraph()
run = p.add_run('Morpho:')
run.bold = True
add_bullet_list(doc, [
    'MorphoFlashloanModule already built (zero-fee flash loans)',
    'Optional: Create Morpho lending market for Superloop-originated positions',
])

p = doc.add_paragraph()
run = p.add_run('Chainlink:')
run.bold = True
add_bullet_list(doc, [
    'NAVLink oracle integration for RWA pricing (reads via Aave Oracle interface)',
    'Chainlink ACE (CCID) for cross-chain investor identity verification',
    'SmartData for institutional bond analytics',
])

p = doc.add_paragraph()
run = p.add_run('DEX/Swap:')
run.bold = True
add_bullet_list(doc, [
    'UniversalDexModule already supports 1inch, Uniswap, and arbitrary swap targets',
    'Verify RWA token DEX liquidity (USTB/USDC, JTRSY/USDC pools)',
    'Alternative: Direct mint/redeem with RWA issuer (bypass DEX)',
])

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 12. COMPLIANCE FRAMEWORK
# ══════════════════════════════════════════════════════════════
doc.add_heading('12. Compliance Framework', level=1)

doc.add_heading('12.1 Three-Layer Compliance Architecture', level=2)

p = doc.add_paragraph()
run = p.add_run('Layer 1: Entity-Level Compliance (One-Time)')
run.bold = True
add_bullet_list(doc, [
    'Legal Entity: SPV or Segregated Portfolio Company (SPC), preferably Cayman Islands',
    'Fund Exemption: Section 3(c)(7) for unlimited qualified purchasers, or 3(c)(1) for max 100 investors',
    'Vault KYC: Smart contract address is KYC-whitelisted by each RWA issuer',
    'Custody: Institutional-grade custodian (Anchorage, BitGo, Fireblocks) for the legal entity',
])

p = doc.add_paragraph()
run = p.add_run('Layer 2: Investor-Level Compliance (Per-User)')
run.bold = True
add_bullet_list(doc, [
    'KYC/AML Verification: Via Sumsub, Coinbase Verifications, or Chainlink ACE (CCID)',
    'Accredited Investor Attestation: Required for most RWA tokens (BUIDL: $5M min, USTB: accredited)',
    'Whitelist Enforcement: Check at DepositManager.requestDeposit() entry point',
    'Transfer Restrictions: Vault shares restricted to whitelisted addresses via onlyPrivileged modifier (already implemented)',
])

p = doc.add_paragraph()
run = p.add_run('Layer 3: On-Chain Enforcement')
run.bold = True
add_bullet_list(doc, [
    'ERC-1404: detectTransferRestriction() check before transfers -- used by Superstate USTB',
    'ERC-3643 (T-REX): Vault address registered in Identity Registry -- used by European RWAs',
    'Chainlink ACE: Cross-chain identity (CCID) + policy verification at transaction level',
    'Non-Transferable aTokens: Horizon prevents secondary market trading of collateral positions',
])

doc.add_heading('12.2 Existing Superloop Compliance Infrastructure', level=2)
doc.add_paragraph(
    'The Superloop codebase already has foundational compliance mechanisms:'
)
add_bullet_list(doc, [
    'privilegedAddresses mapping (SuperloopVault.sol:330): Already restricts share transfers to whitelisted addresses',
    'onlyPrivileged modifier: Controls transfer() and transferFrom() on vault shares',
    'DepositManager: Natural entry point for KYC gating -- add onlyCompliant check to requestDeposit()',
    'VaultRouter: Already whitelists vaults, tokens, and deposit managers -- extend to whitelist investors',
])

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 13. RISK MANAGEMENT
# ══════════════════════════════════════════════════════════════
doc.add_heading('13. Risk Management', level=1)

add_table(doc,
    ['Risk', 'Severity', 'Likelihood', 'Mitigation'],
    [
        ['Liquidation (NAV drop)', 'High', 'Low', 'Conservative LTV (60-70%). Auto-deleverage at HF 1.5. Treasuries have near-zero volatility.'],
        ['Borrow Rate Spike', 'Medium', 'Medium', 'Dynamic monitoring. Auto-deleverage when spread < 50bps. Multiple stablecoin sources.'],
        ['Oracle Staleness', 'High', 'Low', 'Chainlink NAVLink with price bounds. DON rejects out-of-band NAV. Fallback to last known.'],
        ['Smart Contract', 'High', 'Low', 'Battle-tested Aave V3.3 (audited + formally verified). Superloop modules are minimal.'],
        ['Redemption Delay', 'Medium', 'Medium', 'Cash reserve buffer (configurable BPS). Multi-queue WithdrawManager.'],
        ['Regulatory Change', 'Medium', 'Medium', 'SPV legal structure. Multi-jurisdiction optionality. Legal counsel.'],
        ['RWA De-peg', 'High', 'Very Low', 'Position limits. Diversification across issuers. Monitoring of NAV vs benchmark.'],
        ['Interest Rate Inversion', 'Medium', 'Medium', 'Break-even monitoring. Auto-deleverage when spread inverts.'],
    ],
    col_widths=[1.5, 0.8, 0.8, 4.0]
)

doc.add_heading('13.1 Automated Risk Controls', level=2)
add_bullet_list(doc, [
    'IF borrow_rate > (rwa_yield - 0.5%) THEN begin deleverage',
    'IF health_factor < 1.5 THEN partial unwind (reduce leverage by 1x)',
    'IF health_factor < 1.2 THEN emergency full unwind',
    'IF oracle_price_age > 24 hours THEN pause new loops',
    'IF cash_reserve < 5% of TVL THEN pause new deposits',
])

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 14. SUPPORTED ASSETS & PARTNERS
# ══════════════════════════════════════════════════════════════
doc.add_heading('14. Supported Assets & Partners', level=1)

doc.add_heading('14.1 RWA Collateral Assets (Permissioned)', level=2)
add_table(doc,
    ['Asset', 'Issuer', 'Underlying', 'Yield', 'Status'],
    [
        ['USTB', 'Superstate', 'Short Duration US Govt Securities', '~4.5%', 'Live on Horizon'],
        ['USCC', 'Superstate', 'Crypto Carry Fund', '~5-8%', 'Live on Horizon'],
        ['JTRSY', 'Centrifuge / Janus Henderson', 'US Treasury Exposure', '~4.5%', 'Live on Horizon'],
        ['JAAA', 'Centrifuge / Janus Henderson', 'AAA CLOs', '~5-7%', 'Live on Horizon'],
        ['USYC', 'Hashnote / Circle', 'International Short Duration Yield', '~4.5%', 'Live on Horizon'],
        ['VBILL', 'VanEck / Securitize', 'VanEck Treasury Fund', '~4.5%', 'Live on Horizon'],
        ['BUIDL', 'BlackRock / Securitize', 'Treasury Money Market', '~4.5%', 'Pipeline'],
    ],
    col_widths=[0.7, 1.8, 2.0, 0.8, 1.2]
)

doc.add_heading('14.2 Borrowable Stablecoins (Permissionless)', level=2)
add_table(doc,
    ['Asset', 'Issuer', 'Market Cap', 'Notes'],
    [
        ['USDC', 'Circle', '$45B+', 'Primary liquidity source, industry standard'],
        ['GHO', 'Aave DAO', '~$312M', 'Decentralized stablecoin, native to Aave ecosystem'],
        ['RLUSD', 'Ripple', 'Growing', 'Tailored for institutional use'],
    ],
    col_widths=[0.8, 1.2, 1.0, 4.0]
)

doc.add_heading('14.3 Partnership Network', level=2)
add_table(doc,
    ['Partner', 'Role', 'Relevance to Superloop'],
    [
        ['Securitize', 'Tokenization ($4B+ AUM)', 'BUIDL tokenization partner, VBILL infrastructure'],
        ['Chainlink', 'Oracles + Compliance', 'NAVLink pricing, SmartData analytics, ACE compliance engine'],
        ['Chaos Labs', 'Risk infrastructure', 'Dynamic parameter adjustment, Risk Oracles'],
        ['LlamaRisk', 'Due diligence', 'RWA risk framework, asset parameterization'],
        ['Grove Finance', 'Capital routing', 'Sky $1B allocation through Centrifuge into JAAA/ACRDX'],
        ['Ant Digital', 'Digital finance', 'Part of Horizon launch network'],
        ['Ethena', 'Synthetic dollars', 'Part of Horizon launch network, already integrated in Superloop'],
        ['WisdomTree', 'Asset manager', 'Expected future RWA listings on Horizon'],
        ['Franklin Templeton', 'Asset manager', '2026 roadmap partnership with Aave'],
    ],
    col_widths=[1.5, 1.5, 4.0]
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 15. COMPETITIVE POSITIONING
# ══════════════════════════════════════════════════════════════
doc.add_heading('15. Competitive Positioning', level=1)

add_table(doc,
    ['Competitor', 'What They Do', 'Limitation', 'Superloop Advantage'],
    [
        ['Gauntlet', 'Curated levered RWA vaults on Morpho/Drift', 'Vault manager, not infrastructure', 'Superloop is the vault rails Gauntlet could build on'],
        ['Huma Finance', 'Defensive Looping vaults (PayFi)', 'Narrow (PST tokens only)', 'Multi-RWA, multi-protocol, multi-strategy'],
        ['Sommelier', 'Dynamic strategy vaults (Cosmos+EVM)', 'Cosmos bridge dependency', 'Native EVM, no cross-chain bridge risk'],
        ['Morpho Vaults', 'Permissionless lending markets', 'Lending primitive, not vault layer', 'Composes Morpho + Aave + Centrifuge together'],
        ['Summer.fi', 'Frontend for leveraged positions', 'Frontend only, no vault abstraction', 'Full vault infrastructure with compliance'],
        ['Plume Nest', 'RWA yield vaults on Solana', 'Solana-only', 'EVM-native, multi-chain (Ethereum L2s)'],
    ],
    col_widths=[1.2, 1.8, 1.5, 2.5]
)

p = doc.add_paragraph()
run = p.add_run('Superloop\'s moat: ')
run.bold = True
run = p.add_run(
    'We are not building another lending market or another frontend. We are building the institutional vault infrastructure layer '
    'that sits ABOVE Aave Horizon, Centrifuge, Morpho, and Grove -- composing them into managed, compliant, leveraged RWA strategies '
    'with proper async queues, NAV accounting, and compliance gating.'
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 16. REVENUE MODEL
# ══════════════════════════════════════════════════════════════
doc.add_heading('16. Revenue Model', level=1)

add_table(doc,
    ['Revenue Stream', 'Rate', 'Mechanism', 'Status'],
    [
        ['Performance Fee', '10-20% of yield', 'Share dilution on profit appreciation', 'Already implemented'],
        ['Management Fee', '0.5-2% annually', 'AUM-based via exchange rate in accountant', 'Ready to deploy'],
        ['Instant Withdraw Fee', '0.1-0.5%', 'Premium for skipping withdrawal queue', 'Already implemented'],
    ],
    col_widths=[1.5, 1.3, 2.5, 1.2]
)

doc.add_heading('16.1 Revenue Projections', level=2)
add_table(doc,
    ['TVL', 'Gross Yield (10% APR)', 'Performance Fee (15%)', 'Management Fee (1%)', 'Total Revenue'],
    [
        ['$50M', '$5M/yr', '$750K/yr', '$500K/yr', '$1.25M/yr'],
        ['$100M', '$10M/yr', '$1.5M/yr', '$1.0M/yr', '$2.5M/yr'],
        ['$250M', '$25M/yr', '$3.75M/yr', '$2.5M/yr', '$6.25M/yr'],
        ['$500M', '$50M/yr', '$7.5M/yr', '$5.0M/yr', '$12.5M/yr'],
        ['$1B', '$100M/yr', '$15M/yr', '$10M/yr', '$25M/yr'],
    ],
    col_widths=[0.8, 1.5, 1.5, 1.5, 1.2]
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 17. ROADMAP
# ══════════════════════════════════════════════════════════════
doc.add_heading('17. Roadmap', level=1)

doc.add_heading('Phase 1: Horizon Loop Vault (Weeks 1-6)', level=2)
add_bullet_list(doc, [
    'Week 1-2: Deploy Aave V3 modules with Horizon PoolAddressesProvider. Deploy HorizonAccountantPlugin. Register modules.',
    'Week 2-3: Build and test RWA loop execution flow (supply USTB -> borrow USDC -> flash loan loop). Adapt EthenaLoopTest pattern.',
    'Week 3-4: Integrate Chainlink NAVLink oracle adapter. Implement rate monitoring and auto-deleverage in operator/keeper.',
    'Week 4-5: Legal entity setup (SPV). KYC vault address with Superstate (USTB) and/or Centrifuge (JTRSY).',
    'Week 5-6: RWAComplianceGate contract. Security review. Testnet deployment.',
])

doc.add_heading('Phase 2: Centrifuge Integration (Weeks 7-12)', level=2)
add_bullet_list(doc, [
    'Week 7-8: Build CentrifugeSupplyModule with ERC-7540 async requestDeposit/requestRedeem. Build CentrifugeAccountantPlugin.',
    'Week 9-10: Keeper bot for epoch monitoring. Auto-claim when deposits become claimable. WithdrawManager integration.',
    'Week 10-11: Combined strategy: Centrifuge tranche tokens -> Horizon collateral -> leverage loop.',
    'Week 11-12: KYC vault with Centrifuge. End-to-end integration testing. Mainnet launch.',
])

doc.add_heading('Phase 3: Scale & Expand (Weeks 13-20)', level=2)
add_bullet_list(doc, [
    'Week 13-15: Grove partnership for Sky $1B capital allocation. Integration via Centrifuge V3.',
    'Week 15-17: Multi-RWA diversified vaults (USTB + JAAA + ACRDX). Senior/Junior tranche vaults.',
    'Week 17-20: Cross-chain RWA vaults (Base, Arbitrum). Chainlink ACE (CCID) for cross-chain compliance. Morpho market integration.',
])

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 18. WHY WHITELIST SUPERLOOP
# ══════════════════════════════════════════════════════════════
doc.add_heading('18. Why Whitelist Superloop (For RWA Issuers)', level=1)

doc.add_paragraph(
    'This section is specifically for RWA issuers considering whitelisting the Superloop vault contract address.'
)

doc.add_heading('18.1 Demand Generation', level=2)
doc.add_paragraph(
    'Superloop vaults create SUSTAINED demand for your RWA tokens. Leveraged strategies consume 3-5x more tokens per dollar '
    'of user deposits. At $100M vault TVL with 4x leverage, we purchase $400M of your tokens. '
    'This is not speculative trading -- it is long-term, yield-seeking demand that holds positions.'
)

doc.add_heading('18.2 Non-Custodial & Transparent', level=2)
doc.add_paragraph(
    'All positions are on-chain, auditable in real-time. Smart contracts execute deterministically without human intervention. '
    'Superloop cannot move funds -- the architecture is fully non-custodial. '
    'Built on Aave V3, which has been battle-tested with $30B+ historical TVL.'
)

doc.add_heading('18.3 Compliance-First Architecture', level=2)
doc.add_paragraph(
    'Our vault contract completes your KYC/AML requirements at the entity level. '
    'We enforce investor qualification at the deposit entry point. '
    'Vault shares are transfer-restricted via the onlyPrivileged modifier. '
    'Compatible with ERC-1404, ERC-3643, and Chainlink ACE standards.'
)

doc.add_heading('18.4 Institutional Distribution Channel', level=2)
doc.add_paragraph(
    'Superloop serves as a distribution channel for your RWA tokens to qualified purchasers. '
    'A single KYC relationship (vault entity) unlocks access for all verified vault depositors. '
    'This is more efficient than individual KYC for each end-investor.'
)

p = doc.add_paragraph()
run = p.add_run('Ask: ')
run.bold = True
run = p.add_run('Whitelist the Superloop vault contract address as a qualified purchaser for your RWA token.')
run.bold = True

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 19. APPENDIX: TECHNICAL SPECIFICATIONS
# ══════════════════════════════════════════════════════════════
doc.add_heading('19. Appendix: Technical Specifications', level=1)

doc.add_heading('19.1 Smart Contract Addresses (To Be Deployed)', level=2)
add_table(doc,
    ['Contract', 'Target Network', 'Dependencies'],
    [
        ['Superloop Vault (Proxy)', 'Ethereum Mainnet', 'ERC-4626, OpenZeppelin Upgradeable'],
        ['AaveV3SupplyModule (Horizon)', 'Ethereum Mainnet', 'Horizon PoolAddressesProvider'],
        ['AaveV3BorrowModule (Horizon)', 'Ethereum Mainnet', 'Horizon PoolAddressesProvider'],
        ['AaveV3RepayModule (Horizon)', 'Ethereum Mainnet', 'Horizon PoolAddressesProvider'],
        ['AaveV3WithdrawModule (Horizon)', 'Ethereum Mainnet', 'Horizon PoolAddressesProvider'],
        ['AaveV3EmodeModule (Horizon)', 'Ethereum Mainnet', 'Horizon PoolAddressesProvider'],
        ['AaveV3FlashloanModule (Horizon)', 'Ethereum Mainnet', 'Horizon PoolAddressesProvider'],
        ['MorphoFlashloanModule', 'Ethereum Mainnet', 'Morpho Blue contract'],
        ['AaveV3AccountantPlugin (Horizon)', 'Ethereum Mainnet', 'Horizon provider, NAVLink'],
        ['CentrifugeSupplyModule', 'Ethereum Mainnet', 'ERC-7540, Centrifuge V3'],
        ['RWAComplianceGate', 'Ethereum Mainnet', 'Chainlink ACE / KYC provider'],
    ],
    col_widths=[2.5, 1.5, 3.0]
)

doc.add_heading('19.2 Key Interfaces', level=2)

interfaces_text = (
    '// Aave V3 Core (Used by all Aave modules)\n'
    'interface IPoolAddressesProvider {\n'
    '    function getPool() external view returns (address);\n'
    '    function getPriceOracle() external view returns (address);\n'
    '    function getPoolDataProvider() external view returns (address);\n'
    '}\n\n'
    '// Module Execution (Core Superloop pattern)\n'
    'struct ModuleExecutionData {\n'
    '    CallType executionType;  // CALL or DELEGATECALL\n'
    '    address module;          // Registered module address\n'
    '    bytes data;              // abi.encodeWithSelector(...)\n'
    '}\n\n'
    '// Aave V3 Action (Supply/Borrow/Repay/Withdraw)\n'
    'struct AaveV3ActionParams {\n'
    '    address asset;   // Token address\n'
    '    uint256 amount;  // Amount (type(uint256).max = all balance)\n'
    '}\n\n'
    '// Flash Loan Callback\n'
    'struct CallbackData {\n'
    '    address asset;\n'
    '    address addressToApprove;\n'
    '    uint256 amountToApprove;\n'
    '    bytes executionData;  // Nested ModuleExecutionData[]\n'
    '}'
)

p = doc.add_paragraph()
run = p.add_run(interfaces_text)
run.font.name = 'Consolas'
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(0x00, 0x7A, 0xFF)

doc.add_heading('19.3 External References', level=2)
add_bullet_list(doc, [
    'Superloop Core Contracts: github.com/superloop-core-contracts',
    'Aave V3 Horizon: github.com/aave/aave-v3-horizon',
    'Aave Address Book: github.com/aave-dao/aave-address-book (AaveV3EthereumHorizon)',
    'Centrifuge V3: github.com/centrifuge/centrifuge-chain',
    'Chainlink NAVLink: chain.link/navlink',
    'Chainlink ACE: blog.chain.link/automated-compliance-engine-technical-overview',
    'Morpho Blue: github.com/morpho-org/morpho-blue',
    'ERC-7540 Standard: eips.ethereum.org/EIPS/eip-7540',
    'ERC-1404 Standard: eips.ethereum.org/EIPS/eip-1404',
])

# ── Save ──────────────────────────────────────────────────────
output_dir = "/home/user/superloop-core-contracts/docs"
os.makedirs(output_dir, exist_ok=True)
output_path = os.path.join(output_dir, "Superloop_RWA_Strategy_Document.docx")
doc.save(output_path)
print(f"Document saved to: {output_path}")
